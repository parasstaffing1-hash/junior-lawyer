from io import BytesIO

import fitz
import pytest
from docx import Document
from fastapi.testclient import TestClient

from app.main import app
from app.tools.legal_document_parser.models import DocumentFormat, ParseOptions
from app.tools.legal_document_parser.service import LegalDocumentParserError, parse_legal_document

client = TestClient(app)


def _make_pdf() -> bytes:
    document = fitz.open()
    page1 = document.new_page(width=595, height=842)
    page1.insert_text((72, 72), "SERVICE AGREEMENT", fontsize=18)
    page1.insert_text((72, 110), "This agreement starts on 1 August 2026.", fontsize=11)
    page2 = document.new_page(width=595, height=842)
    page2.insert_text((72, 72), "PAYMENT", fontsize=16)
    page2.insert_text((72, 110), "The Client shall pay the invoice within 30 days.", fontsize=11)
    data = document.tobytes()
    document.close()
    return data


def _make_docx() -> bytes:
    document = Document()
    document.core_properties.title = "Sample Contract"
    document.core_properties.author = "Example Lawyer"
    document.add_heading("Services Agreement", level=1)
    document.add_paragraph("This agreement is made between Alpha Ltd and Beta Ltd.")
    document.add_heading("Payment", level=2)
    document.add_paragraph("The Client shall pay within 30 days.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Item"
    table.cell(0, 1).text = "Amount"
    table.cell(1, 0).text = "Fee"
    table.cell(1, 1).text = "1000"
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def test_parse_pdf_preserves_pages_and_text() -> None:
    result = parse_legal_document(_make_pdf(), original_filename="agreement.pdf")
    assert result.metadata.detected_format == DocumentFormat.PDF
    assert result.metadata.page_count == 2
    assert len(result.pages) == 2
    assert "SERVICE AGREEMENT" in result.text
    assert "within 30 days" in result.pages[1].text
    assert result.metadata.sha256


def test_pdf_heading_detection_is_auditable() -> None:
    result = parse_legal_document(_make_pdf())
    assert any(item.text == "SERVICE AGREEMENT" for item in result.headings)
    assert all(item.method.value == "pdf_font_heuristic" for item in result.headings)
    assert any("inferred" in warning.lower() for warning in result.warnings)


def test_parse_docx_preserves_headings_and_tables() -> None:
    result = parse_legal_document(_make_docx(), original_filename="contract.docx")
    assert result.metadata.detected_format == DocumentFormat.DOCX
    assert result.metadata.title == "Sample Contract"
    assert result.metadata.author == "Example Lawyer"
    assert result.metadata.page_count is None
    assert [heading.level for heading in result.headings] == [1, 2]
    assert result.tables[0].rows[1] == ["Fee", "1000"]
    assert result.metadata.table_count == 1
    assert "Alpha Ltd" in result.text


def test_docx_body_order_places_table_after_paragraphs() -> None:
    result = parse_legal_document(_make_docx())
    assert result.text.index("The Client shall pay") < result.text.index("Item\tAmount")
    assert result.blocks[-1].block_type.value == "table"


def test_options_can_suppress_large_response_sections() -> None:
    result = parse_legal_document(
        _make_docx(),
        ParseOptions(include_text=False, include_blocks=False, include_tables=False),
    )
    assert result.text is None
    assert result.blocks == []
    assert result.tables == []
    assert result.metadata.character_count > 0
    assert len(result.headings) == 2


def test_scanned_like_pdf_reports_ocr_warning() -> None:
    document = fitz.open()
    document.new_page()
    data = document.tobytes()
    document.close()
    result = parse_legal_document(data)
    assert result.text == ""
    assert any("ocr" in warning.lower() for warning in result.warnings)


def test_unsupported_file_is_rejected() -> None:
    with pytest.raises(LegalDocumentParserError, match="PDF or DOCX"):
        parse_legal_document(b"plain text is not supported")


def test_character_limit_is_enforced() -> None:
    document = Document()
    document.add_paragraph("A" * 10_500)
    stream = BytesIO()
    document.save(stream)
    with pytest.raises(LegalDocumentParserError, match="character limit"):
        parse_legal_document(stream.getvalue(), ParseOptions(max_extracted_chars=10_000))


def test_parse_api_accepts_pdf() -> None:
    response = client.post(
        "/api/v1/tools/legal-documents/parse",
        files={"file": ("case.pdf", _make_pdf(), "application/pdf")},
        data={"options_json": '{"include_blocks": true}'},
    )
    assert response.status_code == 200
    payload = response.json()
    assert payload["metadata"]["original_filename"] == "case.pdf"
    assert payload["metadata"]["detected_format"] == "pdf"
    assert payload["metadata"]["page_count"] == 2


def test_parse_api_accepts_docx() -> None:
    response = client.post(
        "/api/v1/tools/legal-documents/parse",
        files={
            "file": (
                "case.docx",
                _make_docx(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 200
    payload = response.json()
    assert payload["metadata"]["detected_format"] == "docx"
    assert payload["metadata"]["table_count"] == 1


def test_formats_endpoint_lists_pdf_and_docx() -> None:
    response = client.get("/api/v1/tools/legal-documents/formats")
    assert response.status_code == 200
    assert {item["format"] for item in response.json()} == {"pdf", "docx"}
