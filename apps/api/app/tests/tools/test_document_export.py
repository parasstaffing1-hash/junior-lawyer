import io

import fitz
import pytest
from docx import Document
from fastapi.testclient import TestClient

from app.main import app
from app.tools.document_export.models import DocumentExportRequest
from app.tools.document_export.service import generate_export, preview_export


client = TestClient(app)


def generic_request(fmt: str = "pdf") -> dict:
    return {
        "source_type": "generic",
        "output_format": fmt,
        "source": {
            "title": "Sample Legal Export",
            "subtitle": "Matter ABC-123",
            "sections": [
                {
                    "heading": "Background",
                    "paragraphs": ["This is deterministic legal document export content."],
                    "bullet_items": ["First point", "Second point"],
                },
                {
                    "heading": "Schedule",
                    "tables": [
                        {
                            "headers": ["Date", "Event"],
                            "rows": [["2026-08-08", "Document generated"]],
                        }
                    ],
                },
            ],
            "disclaimer": "Review before legal use.",
        },
        "options": {
            "page_size": "a4",
            "margin_mm": 20,
            "header_text": "CONFIDENTIAL",
        },
    }


def test_preview_counts_sections_tables_and_paragraphs():
    preview = preview_export(DocumentExportRequest.model_validate(generic_request()))
    assert preview.title == "Sample Legal Export"
    assert preview.section_count == 2
    assert preview.table_count == 1
    assert preview.paragraph_count == 3
    assert preview.filename.endswith(".pdf")


def test_generate_pdf_is_valid_and_searchable():
    content, info = generate_export(DocumentExportRequest.model_validate(generic_request("pdf")))
    assert content.startswith(b"%PDF")
    assert info.page_count >= 1
    assert len(info.sha256) == 64
    with fitz.open(stream=content, filetype="pdf") as pdf:
        text = "\n".join(page.get_text() for page in pdf)
    assert "Sample Legal Export" in text
    assert "Document generated" in text
    assert "Review before legal use" in text


def test_generate_docx_is_valid_and_contains_content():
    content, info = generate_export(DocumentExportRequest.model_validate(generic_request("docx")))
    assert content.startswith(b"PK")
    assert info.page_count is None
    doc = Document(io.BytesIO(content))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Sample Legal Export" in text
    assert "Background" in text
    assert "Review before legal use" in text
    assert any(cell.text == "Document generated" for table in doc.tables for row in table.rows for cell in row.cells)


def test_custom_filename_gets_correct_extension():
    payload = generic_request("pdf")
    payload["options"]["filename"] = "Client Export"
    _, info = generate_export(DocumentExportRequest.model_validate(payload))
    assert info.filename == "Client-Export.pdf"


def test_table_width_validation():
    payload = generic_request("pdf")
    payload["source"]["sections"][1]["tables"][0]["rows"] = [["only-one-cell"]]
    response = client.post("/api/v1/tools/document-exports/preview", json=payload)
    assert response.status_code == 422


def test_api_generate_pdf_returns_auditable_headers():
    response = client.post("/api/v1/tools/document-exports/generate", json=generic_request("pdf"))
    assert response.status_code == 200
    assert response.headers["content-type"] == "application/pdf"
    assert len(response.headers["x-document-sha256"]) == 64
    assert int(response.headers["x-document-pages"]) >= 1


def test_formats_endpoint_lists_supported_sources():
    response = client.get("/api/v1/tools/document-exports/formats")
    assert response.status_code == 200
    body = response.json()
    assert body["formats"] == ["docx", "pdf"]
    assert "legal_notice" in body["source_types"]
    assert "client_intake" in body["source_types"]
