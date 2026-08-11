from pathlib import Path

import fitz
from docx import Document as DocxDocument

from app.models.document import ExtractionMethod
from app.services.documents.extractor import extract_document
from app.services.language.normalizer import normalize_document_text


def test_document_normalization_preserves_lines() -> None:
    text = "  IN THE HIGH COURT  \r\n\r\n  ABC v. XYZ  \n   Section 138   "
    normalized = normalize_document_text(text)
    assert normalized == "IN THE HIGH COURT\n\nABC v. XYZ\nSection 138"


def test_native_pdf_extraction_keeps_page_boundaries(tmp_path: Path) -> None:
    path = tmp_path / "order.pdf"
    pdf = fitz.open()
    page1 = pdf.new_page()
    page1.insert_text((72, 72), "IN THE HIGH COURT OF DELHI\nSection 138")
    page2 = pdf.new_page()
    page2.insert_text((72, 72), "ABC PRIVATE LIMITED v. XYZ LIMITED\n12 March 2026")
    pdf.save(path)
    pdf.close()

    result = extract_document(path, ".pdf", allow_ocr=False)

    assert result.extraction_method == ExtractionMethod.NATIVE_PDF
    assert result.ocr_used is False
    assert len(result.pages) == 2
    assert "HIGH COURT" in result.pages[0].text
    assert "ABC PRIVATE" in result.pages[1].text


def test_docx_extraction_reads_paragraphs_and_tables(tmp_path: Path) -> None:
    path = tmp_path / "petition.docx"
    doc = DocxDocument()
    doc.add_paragraph("राम कुमार बनाम श्याम कुमार")
    doc.add_paragraph("धारा 420")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "CNR"
    table.cell(0, 1).text = "DLHC010012342026"
    doc.save(path)

    result = extract_document(path, ".docx")

    assert result.extraction_method == ExtractionMethod.DOCX
    assert result.ocr_used is False
    assert len(result.pages) == 1
    assert "राम कुमार" in result.text
    assert "DLHC010012342026" in result.text
