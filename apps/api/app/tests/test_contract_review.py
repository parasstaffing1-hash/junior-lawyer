from pathlib import Path

from docx import Document

from app.services.contracts.review import classify_clause, extract_review_text, segment_paragraphs, _similarity


def test_clause_classifier_english_and_hindi():
    assert classify_clause("Confidentiality", "The recipient shall keep Confidential Information secret.")[0] == "confidentiality"
    assert classify_clause("गोपनीयता", "पक्ष सभी गोपनीय जानकारी को सुरक्षित रखेगा।")[0] == "confidentiality"
    assert classify_clause("12. Limitation of Liability", "Liability is capped at fees paid.")[0] == "liability"


def test_segment_numbered_contract_clauses():
    rows = [
        ("1. Services", None),
        ("Provider shall perform the services described in Schedule A.", None),
        ("2. Payment", None),
        ("Client shall pay INR 100000 within 15 days.", None),
        ("3. Confidentiality", None),
        ("Each party shall keep Confidential Information secret.", None),
    ]
    clauses = segment_paragraphs(rows)
    assert [item.clause_type for item in clauses] == ["appointment_scope", "fees_payment", "confidentiality"]


def test_extract_docx_contract(tmp_path: Path):
    path = tmp_path / "counterparty.docx"
    doc = Document()
    doc.add_heading("1. Services", level=1)
    doc.add_paragraph("Consultant will provide legal operations support.")
    doc.add_heading("2. Confidentiality", level=1)
    doc.add_paragraph("The parties shall keep all confidential information secret.")
    doc.save(path)
    raw, clauses = extract_review_text(path, ".docx")
    assert "Consultant" in raw
    assert len(clauses) == 2
    assert clauses[1].clause_type == "confidentiality"


def test_similarity_is_deterministic():
    identical = _similarity("Liability is capped at fees paid.", "Liability is capped at fees paid.")
    different = _similarity("Liability is unlimited.", "Liability is capped at fees paid.")
    assert identical == 1.0
    assert identical > different


def test_tracked_change_xml_is_written(tmp_path: Path):
    from zipfile import ZipFile
    from app.models.contract import ContractType
    from app.models.contract_review import CounterpartyContractReview, CounterpartyReviewClause, ClauseDeviationStatus
    from app.services.contracts.review import _build_redline_docx

    review = CounterpartyContractReview(
        title="Test review", contract_type=ContractType.SERVICES, source_format="docx",
        source_filename="x.docx", source_storage_key="x", source_sha256="0" * 64,
        language="en", raw_text="", text_length=0,
    )
    review.clauses = [CounterpartyReviewClause(
        clause_type="liability", heading="Liability", source_text="Liability is unlimited.", position=1,
        classification_confidence=.9, similarity=.2, deviation_status=ClauseDeviationStatus.MODIFIED,
        suggested_body_en="Liability is capped at fees paid.", decision="replace",
    )]
    review.findings = []
    output = tmp_path / "redline.docx"
    changes = _build_redline_docx(review, output)
    assert changes[0]["action"] == "replace"
    with ZipFile(output) as archive:
        xml = archive.read("word/document.xml")
    assert b"<w:del " in xml
    assert b"<w:ins " in xml
    assert b"w:delText" in xml
