from __future__ import annotations

import io
import shutil
from dataclasses import dataclass
from pathlib import Path

import fitz
import pytesseract
from docx import Document as DocxDocument
from PIL import Image, ImageOps

from app.core.config import settings
from app.models.document import ExtractionMethod
from app.services.language.normalizer import normalize_document_text


@dataclass(frozen=True, slots=True)
class ExtractedPage:
    page_number: int
    text: str
    extraction_method: ExtractionMethod
    is_scanned: bool


@dataclass(frozen=True, slots=True)
class ExtractionResult:
    pages: list[ExtractedPage]
    extraction_method: ExtractionMethod
    is_scanned: bool
    ocr_used: bool

    @property
    def text(self) -> str:
        return "\n\n".join(page.text for page in self.pages if page.text)


class DocumentExtractionError(RuntimeError):
    pass


def _meaningful_char_count(text: str) -> int:
    return sum(not char.isspace() for char in text)


def _ocr_available() -> bool:
    return shutil.which("tesseract") is not None


def _ocr_image(image: Image.Image) -> str:
    if not _ocr_available():
        raise DocumentExtractionError(
            "OCR is required for this document but the Tesseract executable is not available"
        )

    image = ImageOps.exif_transpose(image).convert("RGB")
    # Mild preprocessing keeps this deterministic and helps low-contrast scans without
    # damaging normal printed legal documents.
    grayscale = ImageOps.grayscale(image)
    autocontrasted = ImageOps.autocontrast(grayscale)
    text = pytesseract.image_to_string(
        autocontrasted,
        lang=settings.ocr_languages,
        config="--oem 1 --psm 6",
    )
    return normalize_document_text(text)


def _extract_pdf(path: Path, *, allow_ocr: bool) -> ExtractionResult:
    pages: list[ExtractedPage] = []
    methods: set[ExtractionMethod] = set()
    ocr_pages = 0

    try:
        pdf = fitz.open(path)
    except Exception as exc:  # pragma: no cover - PyMuPDF gives varied parse errors
        raise DocumentExtractionError(f"Unable to open PDF: {exc}") from exc

    try:
        for index, page in enumerate(pdf):
            native_text = normalize_document_text(page.get_text("text") or "")
            use_ocr = (
                allow_ocr
                and settings.ocr_enabled
                and _meaningful_char_count(native_text) < settings.pdf_native_text_threshold
            )

            final_text = native_text
            method = ExtractionMethod.NATIVE_PDF
            is_scanned = False

            if use_ocr:
                try:
                    scale = max(settings.ocr_dpi / 72.0, 1.0)
                    pixmap = page.get_pixmap(matrix=fitz.Matrix(scale, scale), alpha=False)
                    image = Image.open(io.BytesIO(pixmap.tobytes("png")))
                    ocr_text = _ocr_image(image)
                    # Keep OCR only when it improves the sparse/native extraction.
                    if _meaningful_char_count(ocr_text) > _meaningful_char_count(native_text):
                        final_text = ocr_text
                        method = ExtractionMethod.OCR
                        is_scanned = True
                        ocr_pages += 1
                except DocumentExtractionError:
                    if not native_text:
                        raise

            methods.add(method)
            pages.append(
                ExtractedPage(
                    page_number=index + 1,
                    text=final_text,
                    extraction_method=method,
                    is_scanned=is_scanned,
                )
            )
    finally:
        pdf.close()

    if not pages:
        raise DocumentExtractionError("PDF contains no pages")

    if methods == {ExtractionMethod.NATIVE_PDF}:
        method = ExtractionMethod.NATIVE_PDF
    elif methods == {ExtractionMethod.OCR}:
        method = ExtractionMethod.OCR
    else:
        method = ExtractionMethod.MIXED_PDF

    return ExtractionResult(
        pages=pages,
        extraction_method=method,
        is_scanned=ocr_pages == len(pages),
        ocr_used=ocr_pages > 0,
    )


def _extract_docx(path: Path) -> ExtractionResult:
    try:
        document = DocxDocument(path)
    except Exception as exc:
        raise DocumentExtractionError(f"Unable to open DOCX: {exc}") from exc

    blocks: list[str] = []
    for paragraph in document.paragraphs:
        value = normalize_document_text(paragraph.text)
        if value:
            blocks.append(value)

    for table in document.tables:
        for row in table.rows:
            cells = [normalize_document_text(cell.text) for cell in row.cells]
            row_text = " | ".join(cell for cell in cells if cell)
            if row_text:
                blocks.append(row_text)

    text = "\n".join(blocks)
    return ExtractionResult(
        pages=[
            ExtractedPage(
                page_number=1,
                text=text,
                extraction_method=ExtractionMethod.DOCX,
                is_scanned=False,
            )
        ],
        extraction_method=ExtractionMethod.DOCX,
        is_scanned=False,
        ocr_used=False,
    )


def _extract_text(path: Path) -> ExtractionResult:
    raw = path.read_bytes()
    text = ""
    for encoding in ("utf-8-sig", "utf-16", "cp1252"):
        try:
            text = raw.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    text = normalize_document_text(text)
    return ExtractionResult(
        pages=[
            ExtractedPage(
                page_number=1,
                text=text,
                extraction_method=ExtractionMethod.TEXT,
                is_scanned=False,
            )
        ],
        extraction_method=ExtractionMethod.TEXT,
        is_scanned=False,
        ocr_used=False,
    )


def _extract_image(path: Path, *, allow_ocr: bool) -> ExtractionResult:
    if not allow_ocr or not settings.ocr_enabled:
        raise DocumentExtractionError("Image documents require OCR, but OCR is disabled")
    try:
        with Image.open(path) as image:
            text = _ocr_image(image)
    except DocumentExtractionError:
        raise
    except Exception as exc:
        raise DocumentExtractionError(f"Unable to open image: {exc}") from exc

    return ExtractionResult(
        pages=[
            ExtractedPage(
                page_number=1,
                text=text,
                extraction_method=ExtractionMethod.IMAGE_OCR,
                is_scanned=True,
            )
        ],
        extraction_method=ExtractionMethod.IMAGE_OCR,
        is_scanned=True,
        ocr_used=True,
    )


def extract_document(path: Path, extension: str, *, allow_ocr: bool = True) -> ExtractionResult:
    extension = extension.casefold()
    if extension == ".pdf":
        return _extract_pdf(path, allow_ocr=allow_ocr)
    if extension == ".docx":
        return _extract_docx(path)
    if extension == ".txt":
        return _extract_text(path)
    if extension in {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"}:
        return _extract_image(path, allow_ocr=allow_ocr)
    raise DocumentExtractionError(f"No extractor registered for {extension}")
