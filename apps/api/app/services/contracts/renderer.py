from __future__ import annotations

import hashlib
import re
from pathlib import Path
from typing import Mapping

from docx import Document as WordDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.core.config import settings
from app.models.contract import Contract, ContractLanguage


PLACEHOLDER_RE = re.compile(r"\{\{\s*([A-Za-z0-9_]+)\s*\}\}")


def _display(value: object) -> str:
    if value is None or value == "":
        return "[TO BE COMPLETED]"
    if isinstance(value, bool):
        return "Yes" if value else "No"
    return str(value)


def render_text(text: str | None, variables: Mapping[str, object]) -> str:
    if not text:
        return ""

    def replace(match: re.Match[str]) -> str:
        return _display(variables.get(match.group(1)))

    return PLACEHOLDER_RE.sub(replace, text)


def build_variables(contract: Contract) -> dict[str, object]:
    values: dict[str, object] = {
        **(contract.questionnaire_json or {}),
        "party_a_name": contract.party_a_name,
        "party_b_name": contract.party_b_name,
        "effective_date": contract.effective_date.isoformat() if contract.effective_date else None,
        "governing_state": contract.governing_state,
        "jurisdiction": contract.jurisdiction,
    }
    # Common defaults keep reusable clauses usable across multiple contract types.
    values.setdefault("scope_description", values.get("role_title") or values.get("purpose") or "[TO BE COMPLETED]")
    values.setdefault("fee_amount", values.get("monthly_salary") or "[TO BE COMPLETED]")
    if contract.contract_type.value == "employment" and values.get("monthly_salary"):
        values.setdefault("payment_schedule", "Monthly, subject to applicable deductions")
    else:
        values.setdefault("payment_schedule", "As mutually agreed in writing")
    values.setdefault("confidentiality_term_months", 36)
    values.setdefault("warranty_days", 0)
    values.setdefault("acceptance_days", 10)
    values.setdefault("notice_days", 30)
    values.setdefault("arbitration_city", values.get("governing_state") or "[TO BE COMPLETED]")
    values.setdefault("dispute_mode", "arbitration")
    return values


def contract_storage_root() -> Path:
    root = settings.storage_root.parent / "contracts"
    root.mkdir(parents=True, exist_ok=True)
    return root


def resolve_contract_storage_key(storage_key: str) -> Path:
    root = contract_storage_root().resolve()
    path = (root / storage_key).resolve()
    if root != path and root not in path.parents:
        raise RuntimeError("Invalid contract storage key")
    return path


def _set_base_font(document: WordDocument) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    for style_name in ("Title", "Heading 1", "Heading 2"):
        style = styles[style_name]
        style.font.name = "Aptos Display" if style_name == "Title" else "Aptos"


def _add_bilingual_paragraph(document: WordDocument, english: str, hindi: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(english)
    run.font.size = Pt(10.5)
    if hindi:
        run2 = p.add_run("\n" + hindi)
        run2.font.name = "Nirmala UI"
        run2.font.size = Pt(10.5)


def generate_docx(contract: Contract, *, version_number: int) -> tuple[str, str, str]:
    variables = build_variables(contract)
    document = WordDocument()
    _set_base_font(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(contract.title.upper())
    title_run.bold = True
    title_run.font.size = Pt(16)

    draft = document.add_paragraph()
    draft.alignment = WD_ALIGN_PARAGRAPH.CENTER
    draft_run = draft.add_run(
        "APPROVED" if contract.status.value == "approved" else "DRAFT — LAWYER REVIEW REQUIRED"
    )
    draft_run.bold = True
    draft_run.font.size = Pt(8.5)

    intro_en = (
        f"This Agreement is made between {contract.party_a_name} (\"Party A\") and "
        f"{contract.party_b_name} (\"Party B\") with effect from "
        f"{_display(variables.get('effective_date'))}."
    )
    intro_hi = (
        f"यह समझौता {contract.party_a_name} (\"पक्ष A\") और {contract.party_b_name} "
        f"(\"पक्ष B\") के बीच {_display(variables.get('effective_date'))} से प्रभावी है।"
    )

    if contract.language == ContractLanguage.ENGLISH:
        document.add_paragraph(intro_en)
    elif contract.language == ContractLanguage.HINDI:
        p = document.add_paragraph(intro_hi)
        for run in p.runs:
            run.font.name = "Nirmala UI"
    else:
        _add_bilingual_paragraph(document, intro_en, intro_hi)

    for index, clause in enumerate(sorted(contract.clauses, key=lambda item: item.position), start=1):
        title_en = f"{index}. {clause.title_en}"
        title_hi = f"{index}. {clause.title_hi or clause.title_en}"
        body_en = render_text(clause.body_en, variables)
        body_hi = render_text(clause.body_hi, variables)

        if contract.language == ContractLanguage.ENGLISH:
            document.add_heading(title_en, level=2)
            document.add_paragraph(body_en)
        elif contract.language == ContractLanguage.HINDI:
            heading = document.add_heading(title_hi, level=2)
            for run in heading.runs:
                run.font.name = "Nirmala UI"
            p = document.add_paragraph(body_hi or body_en)
            for run in p.runs:
                run.font.name = "Nirmala UI"
        else:
            heading = document.add_heading(title_en, level=2)
            for run in heading.runs:
                run.bold = True
            if clause.title_hi:
                hp = document.add_paragraph(title_hi)
                for run in hp.runs:
                    run.font.name = "Nirmala UI"
                    run.bold = True
            _add_bilingual_paragraph(document, body_en, body_hi)

    document.add_paragraph()
    sig = document.add_paragraph()
    sig.add_run("For Party A / पक्ष A\n").bold = True
    sig.add_run(f"{contract.party_a_name}\n\nSignature: ____________________\nDate: ________________________")
    document.add_paragraph()
    sig2 = document.add_paragraph()
    sig2.add_run("For Party B / पक्ष B\n").bold = True
    sig2.add_run(f"{contract.party_b_name}\n\nSignature: ____________________\nDate: ________________________")

    safe_title = re.sub(r"[^A-Za-z0-9._-]+", "-", contract.title).strip("-")[:100] or "contract"
    filename = f"{safe_title}-v{version_number}.docx"
    relative = Path(str(contract.id)) / filename
    path = resolve_contract_storage_key(relative.as_posix())
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)
    digest = hashlib.sha256(path.read_bytes()).hexdigest()
    return filename, relative.as_posix(), digest
