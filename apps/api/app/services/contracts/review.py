from __future__ import annotations

import asyncio
import hashlib
import re
import shutil
from datetime import datetime, timezone
from dataclasses import dataclass
from difflib import SequenceMatcher
from pathlib import Path
from uuid import UUID, uuid4

import fitz
from docx import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from fastapi import HTTPException, UploadFile, status
from sqlalchemy import delete, func, or_, select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.core.config import settings
from app.models.contract import ClauseTemplate, ContractRiskLevel, ContractRiskProfile, ContractType
from app.models.contract_review import (
    ClauseDeviationStatus,
    ContractPlaybook,
    ContractPlaybookRule,
    ContractRedlineVersion,
    ContractReviewStatus,
    CounterpartyContractReview,
    CounterpartyReviewClause,
    CounterpartyReviewFinding,
    PlaybookRequirement,
    RedlineStatus,
    ReviewFindingStatus,
    ReviewSourceFormat,
)
from app.models.matter import Matter
from app.models.security import MatterAccessLevel, OrganizationRole
from app.services.security.context import get_current_actor
from app.services.security.permissions import ROLE_BASE_LEVEL, decide_matter_access, visible_matter_ids
from app.schemas.contract_review import ContractReviewListItem, PlaybookCreate, ReviewStats
from app.services.contracts.catalog import CONTRACT_DEFINITIONS
from app.services.contracts.rules import REQUIRED_CLAUSE_TYPES
from app.services.contracts.service import seed_clause_library
from app.services.documents.extractor import DocumentExtractionError, extract_document
from app.services.documents.storage import discard_staged, sanitize_filename, stage_upload
from app.services.language.detector import detect_language


HEADING_RE = re.compile(r"^\s*(?:\d+(?:\.\d+)*[.)]?|[A-Z][.)]|[IVXLC]+[.)])?\s*([A-Za-z\u0900-\u097F][^\n]{1,90})\s*$")
NUMBERED_HEADING_RE = re.compile(r"^\s*(?:\d+(?:\.\d+)*|[A-Z]|[IVXLC]+)[.)]?\s+\S+")

CLAUSE_KEYWORDS: dict[str, tuple[str, ...]] = {
    "definitions": ("definitions", "meaning of", "परिभाष", "अर्थ"),
    "appointment_scope": ("scope", "services", "appointment", "duties", "deliverables", "कार्य", "सेवाएँ", "दायित्व"),
    "fees_payment": ("fees", "payment", "consideration", "salary", "invoice", "compensation", "भुगतान", "शुल्क", "वेतन", "प्रतिफल"),
    "confidentiality": ("confidential", "non-disclosure", "nda", "गोपनीय", "गोपनीयता"),
    "ip": ("intellectual property", "copyright", "work product", "ownership", "license", "बौद्धिक संपदा", "कॉपीराइट", "स्वामित्व"),
    "data_protection": ("data protection", "personal data", "privacy", "security", "डेटा", "गोपनीयता नीति", "व्यक्तिगत डेटा"),
    "warranty": ("warranty", "warranties", "representation", "प्रतिनिधित्व", "वारंटी", "आश्वासन"),
    "acceptance": ("acceptance", "testing", "acceptance criteria", "स्वीकृति", "परीक्षण"),
    "liability": ("limitation of liability", "liability", "damages", "cap", "दायित्व", "हानि", "क्षतिपूर्ति सीमा"),
    "indemnity": ("indemnity", "indemnify", "hold harmless", "भरपाई", "क्षतिपूर्ति"),
    "term_termination": ("termination", "term", "notice period", "expiry", "समाप्ति", "अवधि", "नोटिस अवधि"),
    "governing_law": ("governing law", "jurisdiction", "applicable law", "शासी कानून", "क्षेत्राधिकार", "लागू कानून"),
    "dispute_resolution": ("dispute", "arbitration", "arbitrator", "mediation", "विवाद", "मध्यस्थता", "पंचाट"),
    "notices": ("notices", "notice", "communication", "सूचना", "नोटिस", "संचार"),
    "force_majeure": ("force majeure", "act of god", "beyond reasonable control", "अप्रत्याशित", "दैवीय घटना"),
    "assignment": ("assignment", "assign", "transfer", "हस्तांतरण", "समनुदेशन"),
    "general": ("entire agreement", "amendment", "waiver", "severability", "सम्पूर्ण समझौता", "संशोधन", "त्याग"),
}


@dataclass(frozen=True, slots=True)
class SegmentedClause:
    position: int
    heading: str | None
    text: str
    clause_type: str
    confidence: float


def _clean(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def _comparison_text(text: str) -> str:
    text = _clean(text).casefold()
    text = re.sub(r"\{\{[^}]+\}\}", " ", text)
    text = re.sub(r"[^a-z0-9\u0900-\u097f ]+", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def classify_clause(heading: str | None, text: str) -> tuple[str, float]:
    heading_norm = _clean(heading or "").casefold()
    body_norm = _clean(text).casefold()
    scores: dict[str, float] = {}
    for clause_type, keywords in CLAUSE_KEYWORDS.items():
        score = 0.0
        for keyword in keywords:
            key = keyword.casefold()
            if key in heading_norm:
                score += 4.0
            if key in body_norm:
                score += 1.0
        if score:
            scores[clause_type] = score
    if not scores:
        return "unknown", 0.15
    clause_type, raw = max(scores.items(), key=lambda item: item[1])
    confidence = min(0.98, 0.48 + raw * 0.08)
    return clause_type, round(confidence, 3)


def _looks_like_heading(text: str, style_name: str | None = None) -> bool:
    value = _clean(text)
    if not value or len(value) > 110:
        return False
    if style_name and style_name.casefold().startswith("heading"):
        return True
    if NUMBERED_HEADING_RE.match(value):
        return True
    letters = [ch for ch in value if ch.isalpha()]
    if len(letters) >= 4 and value.upper() == value and len(value.split()) <= 12:
        return True
    lowered = value.casefold().strip(":")
    return any(lowered == key or lowered.startswith(key + " ") for words in CLAUSE_KEYWORDS.values() for key in words if len(key) >= 4)


def segment_paragraphs(paragraphs: list[tuple[str, str | None]]) -> list[SegmentedClause]:
    sections: list[tuple[str | None, list[str]]] = []
    current_heading: str | None = None
    current: list[str] = []
    preamble: list[str] = []

    def flush() -> None:
        nonlocal current, current_heading
        if current:
            sections.append((current_heading, current))
        current = []

    for raw, style in paragraphs:
        text = _clean(raw)
        if not text:
            continue
        if _looks_like_heading(text, style):
            flush()
            current_heading = text
        else:
            if current_heading is None and not sections:
                preamble.append(text)
            else:
                current.append(text)
    flush()

    if preamble:
        sections.insert(0, ("Preamble", preamble))
    if not sections:
        body = "\n".join(_clean(raw) for raw, _ in paragraphs if _clean(raw))
        if body:
            sections = [(None, [body])]

    output: list[SegmentedClause] = []
    for position, (heading, lines) in enumerate(sections, start=1):
        text = "\n".join(lines).strip()
        if not text and heading:
            text = heading
        clause_type, confidence = classify_clause(heading, text)
        output.append(SegmentedClause(position, heading, text, clause_type, confidence))
    return output


def extract_review_text(path: Path, extension: str) -> tuple[str, list[SegmentedClause]]:
    if extension == ".docx":
        document = DocxDocument(path)
        paragraphs = [(p.text, p.style.name if p.style else None) for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                value = " | ".join(_clean(cell.text) for cell in row.cells if _clean(cell.text))
                if value:
                    paragraphs.append((value, None))
        raw = "\n".join(_clean(text) for text, _ in paragraphs if _clean(text))
        return raw, segment_paragraphs(paragraphs)
    if extension == ".pdf":
        doc = fitz.open(path)
        paragraphs: list[tuple[str, str | None]] = []
        try:
            for page in doc:
                for block in page.get_text("blocks"):
                    text = str(block[4] or "")
                    for line in text.splitlines():
                        if _clean(line):
                            paragraphs.append((line, None))
        finally:
            doc.close()
        raw = "\n".join(text for text, _ in paragraphs)
        if len(_comparison_text(raw)) < 80:
            try:
                extracted = extract_document(path, extension, allow_ocr=True)
                if len(_comparison_text(extracted.text)) > len(_comparison_text(raw)):
                    raw = extracted.text
                    paragraphs = [(line, None) for line in raw.splitlines() if _clean(line)]
            except DocumentExtractionError:
                pass
        return raw, segment_paragraphs(paragraphs)
    if extension == ".txt":
        raw = path.read_text(encoding="utf-8", errors="replace")
        paragraphs = [(line, None) for line in raw.splitlines()]
        return raw, segment_paragraphs(paragraphs)
    raise HTTPException(status_code=415, detail="Contract review supports DOCX, PDF and TXT")


def _review_storage(review_id: UUID, filename: str) -> tuple[Path, str]:
    safe = sanitize_filename(filename)
    relative = Path("contract_reviews") / str(review_id) / safe
    destination = (settings.storage_root / relative).resolve()
    root = settings.storage_root.resolve()
    if root not in destination.parents:
        raise RuntimeError("Invalid review storage path")
    destination.parent.mkdir(parents=True, exist_ok=True)
    return destination, relative.as_posix()


async def seed_default_playbooks(db: AsyncSession) -> int:
    await seed_clause_library(db)
    existing = set((await db.scalars(select(ContractPlaybook.contract_type))).all())
    created = 0
    for contract_type in ContractType:
        if contract_type in existing:
            continue
        playbook = ContractPlaybook(
            name=f"Default {contract_type.value.replace('_', ' ').title()} Playbook",
            owner_label="India default",
            contract_type=contract_type,
            risk_profile=ContractRiskProfile.BALANCED,
            settings_json={"engine": "deterministic_review_v1", "lawyer_review_required": True},
        )
        db.add(playbook)
        await db.flush()
        definition = CONTRACT_DEFINITIONS[contract_type.value]
        required = REQUIRED_CLAUSE_TYPES.get(contract_type.value, set())
        for clause_type in definition["clauses"]:
            db.add(ContractPlaybookRule(
                playbook_id=playbook.id,
                code=f"{contract_type.value}.{clause_type}",
                clause_type=clause_type,
                requirement=PlaybookRequirement.REQUIRED if clause_type in required else PlaybookRequirement.OPTIONAL,
                preferred_variant="balanced",
                risk_level=ContractRiskLevel.HIGH if clause_type in required else ContractRiskLevel.MEDIUM,
                guidance_en=f"Compare the counterparty {clause_type.replace('_', ' ')} language with the approved library position.",
                guidance_hi=f"प्रतिपक्ष की {clause_type.replace('_', ' ')} शर्त की स्वीकृत क्लॉज लाइब्रेरी से तुलना करें।",
                config_json={"modified_threshold": 0.78},
            ))
        created += 1
    if created:
        await db.commit()
    return created


async def list_playbooks(db: AsyncSession) -> list[ContractPlaybook]:
    await seed_default_playbooks(db)
    stmt = select(ContractPlaybook).where(ContractPlaybook.active.is_(True))
    actor = get_current_actor()
    if actor is not None:
        stmt = stmt.where(or_(ContractPlaybook.organization_id.is_(None), ContractPlaybook.organization_id == actor.organization_id))
    stmt = stmt.options(selectinload(ContractPlaybook.rules)).order_by(ContractPlaybook.contract_type)
    return list((await db.scalars(stmt)).unique().all())


async def create_playbook(db: AsyncSession, payload: PlaybookCreate) -> ContractPlaybook:
    actor = get_current_actor()
    if actor is not None and actor.role not in {OrganizationRole.OWNER, OrganizationRole.ADMIN, OrganizationRole.PARTNER}:
        raise HTTPException(status_code=403, detail="Playbook management is not permitted for this role")
    existing_stmt = select(ContractPlaybook).where(
        ContractPlaybook.name == payload.name, ContractPlaybook.contract_type == payload.contract_type
    )
    if actor is not None:
        existing_stmt = existing_stmt.where(ContractPlaybook.organization_id == actor.organization_id)
    else:
        existing_stmt = existing_stmt.where(ContractPlaybook.organization_id.is_(None))
    existing = await db.scalar(existing_stmt)
    if existing is not None:
        raise HTTPException(status_code=409, detail="A playbook with this name and contract type already exists")
    playbook = ContractPlaybook(
        organization_id=actor.organization_id if actor else None,
        created_by_user_id=actor.user_id if actor else None,
        name=payload.name, owner_label=payload.owner_label, contract_type=payload.contract_type,
        risk_profile=payload.risk_profile, active=True,
        settings_json={**payload.settings_json, "lawyer_review_required": True},
    )
    db.add(playbook)
    await db.flush()
    for rule in payload.rules:
        db.add(ContractPlaybookRule(
            playbook_id=playbook.id, code=rule.code, clause_type=rule.clause_type,
            requirement=rule.requirement, preferred_variant=rule.preferred_variant,
            risk_level=rule.risk_level, guidance_en=rule.guidance_en, guidance_hi=rule.guidance_hi,
            config_json=rule.config_json,
        ))
    await db.commit()
    stmt = select(ContractPlaybook).where(ContractPlaybook.id == playbook.id).options(selectinload(ContractPlaybook.rules))
    return (await db.scalars(stmt)).one()


async def _get_default_playbook(db: AsyncSession, contract_type: ContractType) -> ContractPlaybook:
    await seed_default_playbooks(db)
    stmt = select(ContractPlaybook).where(
        ContractPlaybook.contract_type == contract_type, ContractPlaybook.active.is_(True)
    ).options(selectinload(ContractPlaybook.rules)).order_by(ContractPlaybook.created_at)
    playbook = (await db.scalars(stmt)).first()
    if playbook is None:
        raise HTTPException(status_code=500, detail="No contract review playbook available")
    return playbook


async def get_review(db: AsyncSession, review_id: UUID) -> CounterpartyContractReview:
    stmt = select(CounterpartyContractReview).where(CounterpartyContractReview.id == review_id).options(
        selectinload(CounterpartyContractReview.playbook).selectinload(ContractPlaybook.rules),
        selectinload(CounterpartyContractReview.clauses).selectinload(CounterpartyReviewClause.findings),
        selectinload(CounterpartyContractReview.findings),
        selectinload(CounterpartyContractReview.redlines),
    ).execution_options(populate_existing=True)
    row = (await db.execute(stmt)).scalar_one_or_none()
    if row is None:
        raise HTTPException(status_code=404, detail="Contract review not found")
    actor = get_current_actor()
    if actor is not None:
        if row.matter_id:
            decision = await decide_matter_access(db, actor, row.matter_id, required=MatterAccessLevel.VIEW)
            if not decision.allowed:
                raise HTTPException(status_code=403, detail=decision.reason)
        elif row.organization_id != actor.organization_id or ROLE_BASE_LEVEL.get(actor.role) is None:
            raise HTTPException(status_code=403, detail="Contract review access is not permitted")
    return row


async def list_reviews(db: AsyncSession, limit: int = 50) -> list[ContractReviewListItem]:
    stmt = select(CounterpartyContractReview).options(
        selectinload(CounterpartyContractReview.clauses), selectinload(CounterpartyContractReview.findings)
    )
    actor = get_current_actor()
    if actor is not None:
        visible = await visible_matter_ids(db, actor)
        conditions = [CounterpartyContractReview.organization_id == actor.organization_id]
        if visible:
            conditions.append(CounterpartyContractReview.matter_id.in_(visible))
        stmt = stmt.where(or_(*conditions))
    stmt = stmt.order_by(CounterpartyContractReview.updated_at.desc()).limit(limit)
    rows = list((await db.scalars(stmt)).unique().all())
    return [ContractReviewListItem(
        id=row.id, title=row.title, counterparty_name=row.counterparty_name, contract_type=row.contract_type,
        status=row.status, language=row.language, health_score=row.health_score, clause_count=len(row.clauses),
        open_high_risks=sum(1 for finding in row.findings if finding.level == ContractRiskLevel.HIGH and finding.status == ReviewFindingStatus.OPEN),
        source_filename=row.source_filename, updated_at=row.updated_at,
    ) for row in rows]


async def _preferred_template(db: AsyncSession, contract_type: ContractType, clause_type: str, variant: str) -> ClauseTemplate | None:
    await seed_clause_library(db)
    stmt = select(ClauseTemplate).where(
        ClauseTemplate.clause_type == clause_type, ClauseTemplate.active.is_(True)
    ).order_by(ClauseTemplate.version.desc())
    candidates = list((await db.scalars(stmt)).all())
    eligible = [item for item in candidates if contract_type.value in (item.contract_types_json or [])]
    return next((item for item in eligible if item.variant_key == variant), None) or next((item for item in eligible if item.variant_key == "balanced"), None)


def _similarity(a: str, b: str) -> float:
    left, right = _comparison_text(a), _comparison_text(b)
    if not left or not right:
        return 0.0
    seq = SequenceMatcher(None, left, right).ratio()
    left_tokens, right_tokens = set(left.split()), set(right.split())
    jaccard = len(left_tokens & right_tokens) / max(1, len(left_tokens | right_tokens))
    return round((seq * 0.7) + (jaccard * 0.3), 4)


def _deduction(level: ContractRiskLevel) -> int:
    return {ContractRiskLevel.HIGH: 14, ContractRiskLevel.MEDIUM: 7, ContractRiskLevel.LOW: 2}[level]


async def _analyze(db: AsyncSession, review: CounterpartyContractReview, segmented: list[SegmentedClause]) -> None:
    await db.execute(delete(CounterpartyReviewFinding).where(CounterpartyReviewFinding.review_id == review.id))
    await db.execute(delete(CounterpartyReviewClause).where(CounterpartyReviewClause.review_id == review.id))
    await db.flush()
    playbook = review.playbook or await _get_default_playbook(db, review.contract_type)
    review.playbook_id = playbook.id
    rule_map = {rule.clause_type: rule for rule in playbook.rules}
    seen_types: set[str] = set()
    findings: list[CounterpartyReviewFinding] = []

    for seg in segmented:
        rule = rule_map.get(seg.clause_type)
        variant = rule.preferred_variant if rule else playbook.risk_profile.value
        if variant not in {"balanced", "pro_party_a", "pro_party_b"}:
            variant = "balanced"
        template = None
        similarity = 0.0
        deviation = ClauseDeviationStatus.UNKNOWN
        if seg.clause_type != "unknown":
            seen_types.add(seg.clause_type)
            template = await _preferred_template(db, review.contract_type, seg.clause_type, variant)
            if template:
                similarity = _similarity(seg.text, template.body_en)
                deviation = ClauseDeviationStatus.MATCHED if similarity >= 0.78 else ClauseDeviationStatus.MODIFIED

        clause = CounterpartyReviewClause(
            review_id=review.id, clause_type=seg.clause_type, heading=seg.heading, source_text=seg.text,
            position=seg.position, classification_confidence=seg.confidence,
            matched_template_id=template.id if template else None, similarity=similarity,
            deviation_status=deviation,
            suggested_title_en=template.title_en if template else None,
            suggested_title_hi=template.title_hi if template else None,
            suggested_body_en=template.body_en if template else None,
            suggested_body_hi=template.body_hi if template else None,
            metadata_json={"engine": "keyword_classifier_v1"},
        )
        db.add(clause)
        await db.flush()

        if seg.clause_type == "unknown":
            findings.append(CounterpartyReviewFinding(
                review_id=review.id, review_clause_id=clause.id,
                rule_code=f"unknown.{seg.position}", clause_type=None,
                title="Unclassified clause requires review",
                explanation="This section did not confidently map to the deterministic clause taxonomy.",
                recommended_action="Identify the legal/commercial purpose before accepting or redlining this language.",
                level=ContractRiskLevel.LOW, metadata_json={"heading": seg.heading},
            ))
        elif deviation == ClauseDeviationStatus.MODIFIED and template:
            level = ContractRiskLevel.HIGH if similarity < 0.35 else ContractRiskLevel.MEDIUM if similarity < 0.65 else ContractRiskLevel.LOW
            findings.append(CounterpartyReviewFinding(
                review_id=review.id, review_clause_id=clause.id,
                rule_code=f"deviation.{seg.position}.{seg.clause_type}", clause_type=seg.clause_type,
                title=f"{seg.clause_type.replace('_', ' ').title()} differs from playbook",
                explanation=f"The counterparty wording has {round(similarity * 100)}% deterministic similarity to the approved {variant} clause. Similarity is a text signal, not an enforceability conclusion.",
                recommended_action="Compare the commercial/legal effect and replace with the approved clause if the deviation is not acceptable.",
                level=level, metadata_json={"similarity": similarity, "template_code": template.code},
            ))

    for rule in playbook.rules:
        present = rule.clause_type in seen_types
        if rule.requirement == PlaybookRequirement.REQUIRED and not present:
            missing_template = await _preferred_template(
                db, review.contract_type, rule.clause_type, rule.preferred_variant
            )
            findings.append(CounterpartyReviewFinding(
                review_id=review.id, rule_code=f"missing.{rule.clause_type}", clause_type=rule.clause_type,
                title=f"Required {rule.clause_type.replace('_', ' ')} clause not detected",
                explanation="The selected playbook expects this clause, but the deterministic classifier did not detect it in the uploaded contract.",
                recommended_action=rule.guidance_en or "Add the approved playbook clause after lawyer review.",
                level=rule.risk_level, metadata_json={
                    "playbook_rule": rule.code,
                    "suggested_title_en": missing_template.title_en if missing_template else None,
                    "suggested_body_en": missing_template.body_en if missing_template else None,
                    "suggested_title_hi": missing_template.title_hi if missing_template else None,
                    "suggested_body_hi": missing_template.body_hi if missing_template else None,
                },
            ))
        if rule.requirement == PlaybookRequirement.PROHIBITED and present:
            findings.append(CounterpartyReviewFinding(
                review_id=review.id, rule_code=f"prohibited.{rule.clause_type}", clause_type=rule.clause_type,
                title=f"Playbook-prohibited {rule.clause_type.replace('_', ' ')} clause detected",
                explanation="The selected playbook marks this clause type as prohibited.",
                recommended_action="Remove or negotiate this clause before approval.", level=rule.risk_level,
                metadata_json={"playbook_rule": rule.code},
            ))

    for finding in findings:
        db.add(finding)
    review.health_score = max(0, 100 - sum(_deduction(f.level) for f in findings))
    review.status = ContractReviewStatus.ANALYZED
    review.metadata_json = {
        **(review.metadata_json or {}), "analysis_engine": "deterministic_contract_review_v1",
        "lawyer_review_required": True, "clause_taxonomy_version": 1,
    }
    await db.flush()


async def _require_review_work(db: AsyncSession, review: CounterpartyContractReview) -> None:
    actor = get_current_actor()
    if actor is None:
        return
    if review.matter_id:
        decision = await decide_matter_access(db, actor, review.matter_id, required=MatterAccessLevel.WORK)
        if not decision.allowed:
            raise HTTPException(status_code=403, detail=decision.reason)
        return
    if review.organization_id != actor.organization_id or ROLE_BASE_LEVEL.get(actor.role) not in {MatterAccessLevel.WORK, MatterAccessLevel.MANAGE}:
        raise HTTPException(status_code=403, detail="Contract review editing is not permitted")


async def create_review_from_upload(
    db: AsyncSession, *, upload: UploadFile, contract_type: ContractType, title: str,
    counterparty_name: str | None = None, matter_id: UUID | None = None,
    internal_contract_id: UUID | None = None, playbook_id: UUID | None = None,
) -> CounterpartyContractReview:
    actor = get_current_actor()
    if matter_id is not None and await db.get(Matter, matter_id) is None:
        raise HTTPException(status_code=404, detail="Matter not found")
    if actor is not None and matter_id is not None:
        decision = await decide_matter_access(db, actor, matter_id, required=MatterAccessLevel.WORK)
        if not decision.allowed:
            raise HTTPException(status_code=403, detail=decision.reason)
    staged = await stage_upload(upload)
    if staged.extension not in {".docx", ".pdf", ".txt"}:
        discard_staged(staged)
        raise HTTPException(status_code=415, detail="Counterparty contract review supports DOCX, PDF and TXT")
    existing = await db.scalar(select(CounterpartyContractReview).where(CounterpartyContractReview.source_sha256 == staged.sha256))
    if existing:
        discard_staged(staged)
        raise HTTPException(status_code=409, detail={"message": "This exact contract file was already uploaded", "review_id": str(existing.id)})

    playbook = await db.get(ContractPlaybook, playbook_id) if playbook_id else await _get_default_playbook(db, contract_type)
    if playbook is None or playbook.contract_type != contract_type:
        discard_staged(staged)
        raise HTTPException(status_code=400, detail="Playbook does not match contract type")

    review_id = uuid4()
    destination, storage_key = _review_storage(review_id, staged.safe_filename)
    shutil.move(str(staged.path), destination)
    try:
        raw, segmented = await asyncio.to_thread(extract_review_text, destination, staged.extension)
        if not _clean(raw):
            raise HTTPException(status_code=422, detail="No readable text found. OCR scanned PDFs before contract review.")
        language_result = detect_language(raw)
        review = CounterpartyContractReview(
            id=review_id,
            organization_id=actor.organization_id if actor else None,
            created_by_user_id=actor.user_id if actor else None,
            matter_id=matter_id, internal_contract_id=internal_contract_id,
            playbook_id=playbook.id, title=title, counterparty_name=counterparty_name,
            contract_type=contract_type, source_format=ReviewSourceFormat(staged.extension.lstrip(".")),
            source_filename=staged.original_filename, source_storage_key=storage_key, source_sha256=staged.sha256,
            language=language_result.language, raw_text=raw, text_length=len(raw),
            metadata_json={"size_bytes": staged.size_bytes, "original_filename": staged.original_filename},
        )
        db.add(review)
        await db.commit()
        review = await get_review(db, review.id)
        await _analyze(db, review, segmented)
        await db.commit()
        return await get_review(db, review.id)
    except Exception:
        destination.unlink(missing_ok=True)
        raise


async def reanalyze_review(db: AsyncSession, review_id: UUID) -> CounterpartyContractReview:
    review = await get_review(db, review_id)
    await _require_review_work(db, review)
    extension = "." + review.source_format.value
    path = (settings.storage_root / review.source_storage_key).resolve()
    raw, segmented = await asyncio.to_thread(extract_review_text, path, extension)
    review.raw_text = raw
    review.text_length = len(raw)
    await _analyze(db, review, segmented)
    await db.commit()
    return await get_review(db, review_id)


async def update_finding_status(db: AsyncSession, review_id: UUID, finding_id: UUID, new_status: ReviewFindingStatus) -> CounterpartyContractReview:
    review = await get_review(db, review_id)
    await _require_review_work(db, review)
    finding = next((item for item in review.findings if item.id == finding_id), None)
    if finding is None:
        raise HTTPException(status_code=404, detail="Review finding not found")
    finding.status = new_status
    await db.commit()
    return await get_review(db, review_id)


async def update_clause_decision(db: AsyncSession, review_id: UUID, clause_id: UUID, decision: str) -> CounterpartyContractReview:
    review = await get_review(db, review_id)
    await _require_review_work(db, review)
    clause = next((item for item in review.clauses if item.id == clause_id), None)
    if clause is None:
        raise HTTPException(status_code=404, detail="Review clause not found")
    clause.decision = decision
    review.status = ContractReviewStatus.IN_NEGOTIATION
    await db.commit()
    return await get_review(db, review_id)


def _add_tracked_text(paragraph, text: str, *, kind: str, change_id: int) -> None:
    """Append a real WordprocessingML insertion/deletion while keeping visible redline styling."""
    wrapper = OxmlElement("w:ins" if kind == "insert" else "w:del")
    wrapper.set(qn("w:id"), str(change_id))
    wrapper.set(qn("w:author"), "Junior Lawyer deterministic engine")
    wrapper.set(qn("w:date"), datetime.now(timezone.utc).isoformat())

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "176B45" if kind == "insert" else "A61B1B")
    r_pr.append(color)
    if kind == "delete":
        strike = OxmlElement("w:strike")
        strike.set(qn("w:val"), "true")
        r_pr.append(strike)
    else:
        highlight = OxmlElement("w:highlight")
        highlight.set(qn("w:val"), "yellow")
        r_pr.append(highlight)
    run.append(r_pr)

    text_node = OxmlElement("w:t" if kind == "insert" else "w:delText")
    text_node.set(qn("xml:space"), "preserve")
    text_node.text = text
    run.append(text_node)
    wrapper.append(run)
    paragraph._p.append(wrapper)


def _build_redline_docx(review: CounterpartyContractReview, path: Path) -> list[dict]:
    doc = DocxDocument()
    doc.add_heading(f"Negotiation Redline — {review.title}", level=0)
    doc.add_paragraph(
        "DRAFT — LAWYER REVIEW REQUIRED. Word tracked-change markup is generated deterministically; "
        "legal effect and enforceability still require lawyer review."
    )
    changes: list[dict] = []
    change_id = 1
    for clause in sorted(review.clauses, key=lambda item: item.position):
        doc.add_heading(clause.heading or clause.clause_type.replace("_", " ").title(), level=1)
        decision = clause.decision or (
            "replace"
            if clause.deviation_status == ClauseDeviationStatus.MODIFIED and clause.suggested_body_en
            else "keep"
        )
        if decision == "replace" and clause.suggested_body_en:
            p_old = doc.add_paragraph()
            _add_tracked_text(p_old, clause.source_text, kind="delete", change_id=change_id)
            change_id += 1
            p_new = doc.add_paragraph()
            _add_tracked_text(p_new, clause.suggested_body_en, kind="insert", change_id=change_id)
            change_id += 1
            changes.append({
                "clause_id": str(clause.id), "clause_type": clause.clause_type,
                "action": "replace", "similarity": clause.similarity,
            })
        elif decision == "remove":
            p_old = doc.add_paragraph()
            _add_tracked_text(p_old, clause.source_text, kind="delete", change_id=change_id)
            change_id += 1
            changes.append({"clause_id": str(clause.id), "clause_type": clause.clause_type, "action": "remove"})
        else:
            doc.add_paragraph(clause.source_text)
            changes.append({"clause_id": str(clause.id), "clause_type": clause.clause_type, "action": "keep"})

    missing_findings = [
        finding for finding in review.findings
        if finding.rule_code.startswith("missing.") and finding.status == ReviewFindingStatus.OPEN
    ]
    if missing_findings:
        doc.add_heading("Playbook additions", level=1)
    for finding in missing_findings:
        suggested = (finding.metadata_json or {}).get("suggested_body_en")
        if not suggested:
            continue
        title = (finding.metadata_json or {}).get("suggested_title_en") or pretty_clause_type(finding.clause_type)
        doc.add_heading(str(title), level=2)
        paragraph = doc.add_paragraph()
        _add_tracked_text(paragraph, str(suggested), kind="insert", change_id=change_id)
        change_id += 1
        changes.append({"finding_id": str(finding.id), "clause_type": finding.clause_type, "action": "add"})

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return changes


def pretty_clause_type(value: str | None) -> str:
    return (value or "Clause").replace("_", " ").title()


async def generate_redline(db: AsyncSession, review_id: UUID) -> ContractRedlineVersion:
    review = await get_review(db, review_id)
    await _require_review_work(db, review)
    next_version = int(await db.scalar(select(func.max(ContractRedlineVersion.version_number)).where(ContractRedlineVersion.review_id == review.id)) or 0) + 1
    relative = Path("contract_reviews") / str(review.id) / f"redline-v{next_version}.docx"
    path = (settings.storage_root / relative).resolve()
    changes = await asyncio.to_thread(_build_redline_docx, review, path)
    digest = hashlib.sha256(path.read_bytes()).hexdigest()
    version = ContractRedlineVersion(
        review_id=review.id, version_number=next_version, label="Deterministic negotiation redline",
        status=RedlineStatus.GENERATED, changes_json=changes,
        generated_filename=f"{sanitize_filename(review.title)}-redline-v{next_version}.docx",
        generated_storage_key=relative.as_posix(), sha256=digest,
    )
    db.add(version)
    review.status = ContractReviewStatus.IN_NEGOTIATION
    await db.commit()
    await db.refresh(version)
    return version


async def get_redline_path(db: AsyncSession, review_id: UUID, redline_id: UUID) -> tuple[ContractRedlineVersion, Path]:
    review = await get_review(db, review_id)
    actor = get_current_actor()
    if actor is not None:
        if review.matter_id:
            decision = await decide_matter_access(db, actor, review.matter_id, required=MatterAccessLevel.VIEW)
            if not decision.allowed or not decision.export_allowed:
                raise HTTPException(status_code=403, detail="Redline export is blocked by security policy")
        else:
            from app.services.security.service import get_policy
            policy = await get_policy(db, actor.organization_id)
            if review.organization_id != actor.organization_id or not policy.allow_exports_default:
                raise HTTPException(status_code=403, detail="Redline export is blocked by security policy")
    version = next((item for item in review.redlines if item.id == redline_id), None)
    if version is None:
        raise HTTPException(status_code=404, detail="Redline version not found")
    path = (settings.storage_root / version.generated_storage_key).resolve()
    if not path.exists():
        raise HTTPException(status_code=404, detail="Generated redline file is missing")
    return version, path


async def get_stats(db: AsyncSession) -> ReviewStats:
    actor = get_current_actor()
    review_filter = []
    review_ids = None
    if actor is not None:
        visible = await visible_matter_ids(db, actor)
        conditions = [CounterpartyContractReview.organization_id == actor.organization_id]
        if visible:
            conditions.append(CounterpartyContractReview.matter_id.in_(visible))
        review_ids = select(CounterpartyContractReview.id).where(or_(*conditions))
        review_filter = [CounterpartyReviewClause.review_id.in_(review_ids)]
    reviews_stmt = select(func.count(CounterpartyContractReview.id))
    if review_ids is not None:
        reviews_stmt = reviews_stmt.where(CounterpartyContractReview.id.in_(review_ids))
    reviews = int(await db.scalar(reviews_stmt) or 0)
    clauses_stmt = select(func.count(CounterpartyReviewClause.id))
    if review_filter:
        clauses_stmt = clauses_stmt.where(*review_filter)
    clauses = int(await db.scalar(clauses_stmt) or 0)
    high_stmt = select(func.count(CounterpartyReviewFinding.id)).where(
        CounterpartyReviewFinding.level == ContractRiskLevel.HIGH,
        CounterpartyReviewFinding.status == ReviewFindingStatus.OPEN,
    )
    redline_stmt = select(func.count(ContractRedlineVersion.id))
    if review_ids is not None:
        high_stmt = high_stmt.where(CounterpartyReviewFinding.review_id.in_(review_ids))
        redline_stmt = redline_stmt.where(ContractRedlineVersion.review_id.in_(review_ids))
    open_high = int(await db.scalar(high_stmt) or 0)
    redlines = int(await db.scalar(redline_stmt) or 0)
    return ReviewStats(reviews=reviews, clauses=clauses, open_high_risks=open_high, redlines=redlines)
