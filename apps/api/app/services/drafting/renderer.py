from __future__ import annotations

import hashlib
import re
from pathlib import Path

from docx import Document as WordDocument
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from app.core.config import settings
from app.models.drafting import LegalDraft, LegalDraftLanguage, LegalDraftStatus


def draft_storage_root() -> Path:
    root = settings.storage_root.parent / "legal_drafts"
    root.mkdir(parents=True, exist_ok=True)
    return root


def resolve_draft_storage_key(storage_key: str) -> Path:
    root = draft_storage_root().resolve()
    path = (root / storage_key).resolve()
    if root != path and root not in path.parents:
        raise RuntimeError("Invalid legal draft storage key")
    return path


def _base_style(document: WordDocument) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for name in ("Title", "Heading 1", "Heading 2"):
        style = document.styles[name]
        style.font.name = "Aptos Display" if name == "Title" else "Aptos"


def _add_text(document: WordDocument, text: str, *, hindi: bool = False) -> None:
    chunks = text.splitlines() or [""]
    for chunk in chunks:
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        run = p.add_run(chunk)
        if hindi:
            run.font.name = "Nirmala UI"
        run.font.size = Pt(11)


def _add_heading(document: WordDocument, text: str, *, hindi: bool = False) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11.5)
    if hindi:
        run.font.name = "Nirmala UI"


def generate_docx(draft: LegalDraft, *, version_number: int) -> tuple[str, str, str]:
    document = WordDocument()
    _base_style(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(draft.title.upper())
    run.bold = True
    run.font.size = Pt(15)

    status = document.add_paragraph()
    status.alignment = WD_ALIGN_PARAGRAPH.CENTER
    status_run = status.add_run(
        "APPROVED BY LAWYER" if draft.status == LegalDraftStatus.APPROVED
        else "DRAFT — LAWYER REVIEW REQUIRED"
    )
    status_run.bold = True
    status_run.font.size = Pt(8.5)

    meta_parts = [part for part in [draft.court_name, draft.case_number] if part]
    if meta_parts:
        meta = document.add_paragraph(" · ".join(meta_parts))
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in meta.runs:
            r.font.size = Pt(9)

    for section in sorted(draft.sections, key=lambda item: item.position):
        if draft.language == LegalDraftLanguage.ENGLISH:
            _add_heading(document, section.title_en)
            _add_text(document, section.body_en)
        elif draft.language == LegalDraftLanguage.HINDI:
            _add_heading(document, section.title_hi or section.title_en, hindi=True)
            _add_text(document, section.body_hi or section.body_en, hindi=True)
        else:
            _add_heading(document, section.title_en)
            if section.title_hi:
                _add_heading(document, section.title_hi, hindi=True)
            _add_text(document, section.body_en)
            if section.body_hi:
                _add_text(document, section.body_hi, hindi=True)

    document.add_section(WD_SECTION.NEW_PAGE)
    review = document.add_paragraph()
    review.alignment = WD_ALIGN_PARAGRAPH.CENTER
    review_run = review.add_run("INTERNAL REVIEW RECORD / आंतरिक समीक्षा अभिलेख")
    review_run.bold = True
    review_run.font.size = Pt(12)
    document.add_paragraph(
        "This page is an internal review aid. Remove it before filing/service if the firm's workflow requires a clean filing copy."
    )
    p_hi = document.add_paragraph(
        "यह पृष्ठ आंतरिक समीक्षा हेतु है। यदि फर्म की प्रक्रिया में स्वच्छ दाखिल प्रति आवश्यक हो तो दाखिल/सेवा से पूर्व इसे हटाएँ।"
    )
    for r in p_hi.runs:
        r.font.name = "Nirmala UI"

    for section in sorted(draft.sections, key=lambda item: item.position):
        marker = "✓" if section.reviewed else "○"
        p = document.add_paragraph(f"{marker} {section.position}. {section.title_en}")
        if section.sources:
            for source in section.sources[:8]:
                locator = f" — {source.locator}" if source.locator else ""
                source_p = document.add_paragraph(
                    f"    Source: {source.label}{locator}"
                )
                for r in source_p.runs:
                    r.font.size = Pt(8.5)

    safe_title = re.sub(r"[^A-Za-z0-9._-]+", "-", draft.title).strip("-")[:100] or "legal-draft"
    filename = f"{safe_title}-v{version_number}.docx"
    relative = Path(str(draft.id)) / filename
    path = resolve_draft_storage_key(relative.as_posix())
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)
    digest = hashlib.sha256(path.read_bytes()).hexdigest()
    return filename, relative.as_posix(), digest
