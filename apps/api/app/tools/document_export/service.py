from __future__ import annotations

import hashlib
import io
import re
from datetime import datetime, timezone
from typing import Any

import fitz
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4, LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)
from xml.sax.saxutils import escape

from app.tools.affidavit.models import AffidavitGenerationResponse
from app.tools.case_timeline.models import CaseTimelineResponse
from app.tools.client_matter_intake.models import ClientMatterIntakeResponse
from app.tools.evidence_index.models import EvidenceIndexResponse
from app.tools.legal_checklist.models import LegalChecklistResponse
from app.tools.legal_notice.models import LegalNoticeGenerationResponse
from app.tools.document_export.models import (
    DocumentExportPreview,
    DocumentExportRequest,
    ExportFormat,
    ExportOptions,
    ExportSourceType,
    ExportTable,
    GeneratedDocument,
    GenericExportDocument,
    GenericExportSection,
    PageSize,
)


MAX_SOURCE_BYTES = 2_000_000
MAX_OUTPUT_BYTES = 25_000_000
DISCLAIMER_HEADING = "Important notice"


class DocumentExportError(ValueError):
    pass


class DocumentExportInputError(DocumentExportError):
    pass


class DocumentExportGenerationError(DocumentExportError):
    pass


def _stringify(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, bool):
        return "Yes" if value else "No"
    if isinstance(value, (list, tuple)):
        return ", ".join(_stringify(item) for item in value)
    if isinstance(value, dict):
        return "; ".join(f"{key}: {_stringify(val)}" for key, val in value.items())
    return str(value)


def _humanize(key: str) -> str:
    return key.replace("_", " ").strip().title()


def _safe_filename(value: str, extension: str) -> str:
    stem = re.sub(r"[^A-Za-z0-9._ -]+", "", value).strip().replace(" ", "-")
    stem = re.sub(r"-+", "-", stem).strip("-._") or "legal-document"
    stem = stem[:120]
    return f"{stem}.{extension}"


def _metadata_section(metadata: dict[str, str]) -> GenericExportSection | None:
    rows = [[_humanize(key), value] for key, value in metadata.items() if value]
    if not rows:
        return None
    return GenericExportSection(
        heading="Document details",
        tables=[ExportTable(headers=["Field", "Value"], rows=rows)],
    )


def _adapt_notice(source: dict[str, Any]) -> GenericExportDocument:
    model = LegalNoticeGenerationResponse.model_validate(source)
    sections: list[GenericExportSection] = []
    meta = _metadata_section(
        {
            "notice_type": model.notice_type,
            "jurisdiction": model.jurisdiction,
            "generation_date": model.generation_date.isoformat(),
            "template": f"{model.template_id} v{model.template_version}",
            "subject": model.subject,
        }
    )
    if meta:
        sections.append(meta)
    for item in model.sections:
        sections.append(GenericExportSection(heading=item.heading, paragraphs=[item.body]))
    if model.warnings:
        sections.append(GenericExportSection(heading="Warnings", bullet_items=model.warnings))
    return GenericExportDocument(
        title=model.title,
        subtitle=model.subject,
        sections=sections,
        disclaimer=model.disclaimer,
    )


def _adapt_affidavit(source: dict[str, Any]) -> GenericExportDocument:
    model = AffidavitGenerationResponse.model_validate(source)
    sections: list[GenericExportSection] = []
    meta = _metadata_section(
        {
            "affidavit_type": model.affidavit_type,
            "jurisdiction": model.jurisdiction,
            "generation_date": model.generation_date.isoformat(),
            "template": f"{model.template_id} v{model.template_version}",
        }
    )
    if meta:
        sections.append(meta)
    for item in model.sections:
        # Avoid duplicating the separately structured statement/annexure lists.
        body = item.body
        if item.id not in {"statements", "annexures"}:
            sections.append(GenericExportSection(heading=item.heading, paragraphs=[body]))
    sections.append(
        GenericExportSection(
            heading="Statements",
            numbered_items=[
                statement.text
                + (f" [Reference: {statement.source_reference}]" if statement.source_reference else "")
                for statement in model.statements
            ],
        )
    )
    if model.annexures:
        sections.append(
            GenericExportSection(
                heading="Annexures",
                tables=[
                    ExportTable(
                        headers=["Label", "Title", "Date", "Description"],
                        rows=[
                            [
                                item.label,
                                item.title,
                                item.document_date.isoformat() if item.document_date else "",
                                item.description or "",
                            ]
                            for item in model.annexures
                        ],
                    )
                ],
            )
        )
    if model.warnings:
        sections.append(GenericExportSection(heading="Warnings", bullet_items=model.warnings))
    return GenericExportDocument(title=model.title, sections=sections, disclaimer=model.disclaimer)


def _adapt_timeline(source: dict[str, Any]) -> GenericExportDocument:
    model = CaseTimelineResponse.model_validate(source)
    rows: list[list[str]] = []
    for event in model.events:
        refs = "; ".join(
            f"{ref.label}{' p. ' + ref.page if ref.page else ''}" for ref in event.source_references
        )
        rows.append(
            [
                str(event.sequence),
                event.display_date,
                event.title,
                event.event_type.value,
                event.importance.value,
                refs,
            ]
        )
    sections: list[GenericExportSection] = []
    meta = _metadata_section(
        {
            "case_reference": model.case_reference or "",
            "event_count": str(model.summary.event_count),
            "first_date": model.summary.first_date.isoformat(),
            "last_date": model.summary.last_date.isoformat(),
            "span_days": str(model.summary.span_days),
        }
    )
    if meta:
        sections.append(meta)
    sections.append(
        GenericExportSection(
            heading="Chronology",
            tables=[ExportTable(headers=["#", "Date", "Event", "Type", "Importance", "Sources"], rows=rows)],
        )
    )
    if model.warnings:
        sections.append(GenericExportSection(heading="Warnings", bullet_items=model.warnings))
    return GenericExportDocument(title=model.title, sections=sections, disclaimer=model.disclaimer)


def _adapt_evidence_index(source: dict[str, Any]) -> GenericExportDocument:
    model = EvidenceIndexResponse.model_validate(source)
    rows = [
        [
            item.label,
            item.document_date.isoformat() if item.document_date else "",
            item.title,
            item.category or "",
            item.page_range or "",
            "Yes" if item.confidential else "No",
        ]
        for item in model.documents
    ]
    sections: list[GenericExportSection] = []
    meta = _metadata_section(
        {
            "case_reference": model.case_reference or "",
            "index_type": model.index_type.value,
            "document_count": str(model.summary.document_count),
            "total_pages": _stringify(model.summary.total_pages),
        }
    )
    if meta:
        sections.append(meta)
    sections.append(
        GenericExportSection(
            heading="Index",
            tables=[ExportTable(headers=["Label", "Date", "Document", "Category", "Pages", "Confidential"], rows=rows)],
        )
    )
    if model.warnings:
        sections.append(GenericExportSection(heading="Warnings", bullet_items=model.warnings))
    return GenericExportDocument(title=model.title, sections=sections, disclaimer=model.disclaimer)


def _adapt_checklist(source: dict[str, Any]) -> GenericExportDocument:
    model = LegalChecklistResponse.model_validate(source)
    rows = [
        [
            item.category,
            item.title,
            item.requirement.value,
            item.status.value,
            "Yes" if item.applicable else "No",
        ]
        for item in model.items
    ]
    sections: list[GenericExportSection] = []
    meta = _metadata_section(
        {
            "matter_type": model.matter_type,
            "jurisdiction": model.jurisdiction,
            "assessment_date": model.assessment_date.isoformat(),
            "template": f"{model.template_id} v{model.template_version}",
            "completion_percent": f"{model.summary.completion_percent:.1f}%",
            "required_completion_percent": f"{model.summary.required_completion_percent:.1f}%",
        }
    )
    if meta:
        sections.append(meta)
    sections.append(
        GenericExportSection(
            heading="Checklist",
            tables=[ExportTable(headers=["Category", "Item", "Requirement", "Status", "Applicable"], rows=rows)],
        )
    )
    if model.summary.outstanding_required_keys:
        sections.append(
            GenericExportSection(
                heading="Outstanding required items",
                bullet_items=model.summary.outstanding_required_keys,
            )
        )
    if model.warnings:
        sections.append(GenericExportSection(heading="Warnings", bullet_items=model.warnings))
    return GenericExportDocument(title=model.title, sections=sections, disclaimer=model.disclaimer)


def _adapt_intake(source: dict[str, Any]) -> GenericExportDocument:
    model = ClientMatterIntakeResponse.model_validate(source)
    rows = []
    for item in model.fields:
        if not item.applicable:
            continue
        rows.append(
            [
                item.label,
                _stringify(item.normalized_value),
                "Yes" if item.required else "No",
                "Yes" if item.valid else "No",
            ]
        )
    sections: list[GenericExportSection] = []
    meta = _metadata_section(
        {
            "matter_type": model.matter_type,
            "client_type": model.client_type,
            "jurisdiction": model.jurisdiction,
            "intake_date": model.intake_date.isoformat(),
            "template": f"{model.template_id} v{model.template_version}",
            "completion_percent": f"{model.summary.completion_percent:.1f}%",
            "ready_for_review": "Yes" if model.summary.ready_for_review else "No",
            "audit_hash_sha256": model.audit_hash_sha256,
        }
    )
    if meta:
        sections.append(meta)
    sections.append(
        GenericExportSection(
            heading="Intake fields",
            tables=[ExportTable(headers=["Field", "Value", "Required", "Valid"], rows=rows)],
        )
    )
    if model.conflict_parties:
        sections.append(
            GenericExportSection(
                heading="Conflict-check inputs",
                tables=[
                    ExportTable(
                        headers=["Name", "Role", "Organization", "Aliases"],
                        rows=[
                            [
                                party.name,
                                party.role.value,
                                party.organization or "",
                                ", ".join(party.aliases),
                            ]
                            for party in model.conflict_parties
                        ],
                    )
                ],
            )
        )
    if model.consents:
        sections.append(
            GenericExportSection(
                heading="Consents",
                tables=[
                    ExportTable(
                        headers=["Consent", "Required", "Accepted", "Accepted at"],
                        rows=[
                            [
                                consent.label,
                                "Yes" if consent.required else "No",
                                "Yes" if consent.accepted else "No",
                                consent.accepted_at.isoformat() if consent.accepted_at else "",
                            ]
                            for consent in model.consents
                            if consent.applicable
                        ],
                    )
                ],
            )
        )
    if model.warnings:
        sections.append(GenericExportSection(heading="Warnings", bullet_items=model.warnings))
    return GenericExportDocument(title=model.title, sections=sections, disclaimer=model.disclaimer)


def _adapt_generic(source: dict[str, Any]) -> GenericExportDocument:
    return GenericExportDocument.model_validate(source)


ADAPTERS = {
    ExportSourceType.LEGAL_NOTICE: _adapt_notice,
    ExportSourceType.AFFIDAVIT: _adapt_affidavit,
    ExportSourceType.CASE_TIMELINE: _adapt_timeline,
    ExportSourceType.EVIDENCE_INDEX: _adapt_evidence_index,
    ExportSourceType.LEGAL_CHECKLIST: _adapt_checklist,
    ExportSourceType.CLIENT_INTAKE: _adapt_intake,
    ExportSourceType.GENERIC: _adapt_generic,
}


def _canonicalize(request: DocumentExportRequest) -> tuple[GenericExportDocument, list[str]]:
    approximate_size = len(repr(request.source).encode("utf-8", errors="ignore"))
    if approximate_size > MAX_SOURCE_BYTES:
        raise DocumentExportInputError(
            f"source payload exceeds the {MAX_SOURCE_BYTES:,}-byte safety limit"
        )
    try:
        document = ADAPTERS[request.source_type](request.source)
    except Exception as exc:
        if isinstance(exc, DocumentExportError):
            raise
        raise DocumentExportInputError(
            f"source payload is not valid for '{request.source_type.value}': {exc}"
        ) from exc

    warnings: list[str] = []
    if request.options.include_disclaimer and not document.disclaimer:
        warnings.append("No disclaimer was supplied by the source document.")
    return document, warnings


def preview_export(request: DocumentExportRequest) -> DocumentExportPreview:
    document, warnings = _canonicalize(request)
    extension = request.output_format.value
    filename = request.options.filename or _safe_filename(document.title, extension)
    if not filename.lower().endswith(f".{extension}"):
        filename = _safe_filename(filename, extension)
    table_count = sum(len(section.tables) for section in document.sections)
    paragraph_count = sum(
        len(section.paragraphs) + len(section.bullet_items) + len(section.numbered_items)
        for section in document.sections
    )
    return DocumentExportPreview(
        source_type=request.source_type,
        output_format=request.output_format,
        title=document.title,
        filename=filename,
        section_count=len(document.sections),
        table_count=table_count,
        paragraph_count=paragraph_count,
        page_size=request.options.page_size,
        warnings=warnings,
    )


def _set_docx_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _configure_docx(document: Document, options: ExportOptions) -> None:
    section = document.sections[0]
    if options.page_size == PageSize.A4:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
    else:
        section.page_width = Mm(215.9)
        section.page_height = Mm(279.4)
    margin = Mm(options.margin_mm)
    section.top_margin = margin
    section.bottom_margin = margin
    section.left_margin = margin
    section.right_margin = margin

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    styles["Title"].font.name = "Times New Roman"
    styles["Title"].font.size = Pt(18)
    styles["Heading 1"].font.name = "Times New Roman"
    styles["Heading 1"].font.size = Pt(14)
    styles["Heading 2"].font.name = "Times New Roman"
    styles["Heading 2"].font.size = Pt(12)

    if options.header_text:
        paragraph = section.header.paragraphs[0]
        paragraph.text = options.header_text
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.runs[0].font.size = Pt(8)
    if options.footer_text or options.include_generated_footer:
        footer_bits = []
        if options.footer_text:
            footer_bits.append(options.footer_text)
        if options.include_generated_footer:
            footer_bits.append("Generated by Lawyer Tools - deterministic export")
        paragraph = section.footer.paragraphs[0]
        paragraph.text = " | ".join(footer_bits)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.runs[0].font.size = Pt(8)


def _add_docx_table(document: Document, table_data: ExportTable) -> None:
    if table_data.title:
        p = document.add_paragraph()
        run = p.add_run(table_data.title)
        run.bold = True
    table = document.add_table(rows=1, cols=len(table_data.headers))
    table.style = "Table Grid"
    header = table.rows[0].cells
    for index, value in enumerate(table_data.headers):
        header[index].text = value
        header[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _set_docx_cell_shading(header[index], "E7E6E6")
        for run in header[index].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
    for row_values in table_data.rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            cells[index].text = value
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for paragraph in cells[index].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    document.add_paragraph()


def _generate_docx(document_model: GenericExportDocument, options: ExportOptions) -> bytes:
    document = Document()
    _configure_docx(document, options)
    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(document_model.title)
    if document_model.subtitle:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(document_model.subtitle)
        run.italic = True

    for section in document_model.sections:
        if section.heading:
            document.add_heading(section.heading, level=1)
        for text in section.paragraphs:
            document.add_paragraph(text)
        for item in section.bullet_items:
            document.add_paragraph(item, style="List Bullet")
        for item in section.numbered_items:
            document.add_paragraph(item, style="List Number")
        for table in section.tables:
            _add_docx_table(document, table)

    if options.include_disclaimer and document_model.disclaimer:
        document.add_heading(DISCLAIMER_HEADING, level=1)
        p = document.add_paragraph(document_model.disclaimer)
        for run in p.runs:
            run.italic = True
            run.font.size = Pt(9)

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _pdf_safe_text(value: str) -> str:
    replacements = {
        "\u2018": "'",
        "\u2019": "'",
        "\u201c": '"',
        "\u201d": '"',
        "\u2013": "-",
        "\u2014": "-",
        "\u2026": "...",
        "\u20b9": "INR ",
        "\u00a0": " ",
    }
    for old, new in replacements.items():
        value = value.replace(old, new)
    return value.encode("latin-1", errors="replace").decode("latin-1")


def _pdf_para(value: str, style) -> Paragraph:
    safe = escape(_pdf_safe_text(value)).replace("\n", "<br/>")
    return Paragraph(safe, style)


def _generate_pdf(document_model: GenericExportDocument, options: ExportOptions) -> bytes:
    output = io.BytesIO()
    page_size = A4 if options.page_size == PageSize.A4 else LETTER
    margin = options.margin_mm * mm

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "LegalTitle",
        parent=styles["Title"],
        fontName="Times-Bold",
        fontSize=17,
        leading=21,
        alignment=TA_CENTER,
        spaceAfter=8,
    )
    subtitle_style = ParagraphStyle(
        "LegalSubtitle",
        parent=styles["Normal"],
        fontName="Times-Italic",
        fontSize=10.5,
        leading=14,
        alignment=TA_CENTER,
        spaceAfter=12,
    )
    heading_style = ParagraphStyle(
        "LegalHeading",
        parent=styles["Heading1"],
        fontName="Times-Bold",
        fontSize=12.5,
        leading=15,
        spaceBefore=9,
        spaceAfter=5,
    )
    body_style = ParagraphStyle(
        "LegalBody",
        parent=styles["BodyText"],
        fontName="Times-Roman",
        fontSize=10.5,
        leading=14,
        alignment=TA_LEFT,
        spaceAfter=5,
    )
    small_style = ParagraphStyle(
        "LegalSmall",
        parent=body_style,
        fontSize=8.5,
        leading=11,
    )

    def header_footer(canvas, doc):
        canvas.saveState()
        width, height = page_size
        canvas.setFont("Helvetica", 7.5)
        if options.header_text:
            canvas.drawCentredString(width / 2, height - 10 * mm, _pdf_safe_text(options.header_text))
        footer_bits = []
        if options.footer_text:
            footer_bits.append(options.footer_text)
        if options.include_generated_footer:
            footer_bits.append("Generated by Lawyer Tools - deterministic export")
        if footer_bits:
            canvas.drawCentredString(width / 2, 8 * mm, _pdf_safe_text(" | ".join(footer_bits)))
        canvas.setFont("Helvetica", 7.5)
        canvas.drawRightString(width - margin, 8 * mm, f"Page {doc.page}")
        canvas.restoreState()

    doc = BaseDocTemplate(
        output,
        pagesize=page_size,
        leftMargin=margin,
        rightMargin=margin,
        topMargin=max(margin, 17 * mm if options.header_text else margin),
        bottomMargin=max(margin, 15 * mm),
        title=_pdf_safe_text(document_model.title),
        author="Lawyer Tools",
    )
    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="normal")
    doc.addPageTemplates([PageTemplate(id="legal", frames=[frame], onPage=header_footer)])

    story = [_pdf_para(document_model.title, title_style)]
    if document_model.subtitle:
        story.append(_pdf_para(document_model.subtitle, subtitle_style))

    for section in document_model.sections:
        if section.heading:
            story.append(_pdf_para(section.heading, heading_style))
        for text in section.paragraphs:
            story.append(_pdf_para(text, body_style))
        for item in section.bullet_items:
            story.append(_pdf_para(f"- {item}", body_style))
        for index, item in enumerate(section.numbered_items, start=1):
            story.append(_pdf_para(f"{index}. {item}", body_style))
        for table_data in section.tables:
            if table_data.title:
                story.append(_pdf_para(table_data.title, body_style))
            data = [[_pdf_para(cell, small_style) for cell in table_data.headers]]
            data.extend([[_pdf_para(cell, small_style) for cell in row] for row in table_data.rows])
            column_width = doc.width / max(1, len(table_data.headers))
            table = Table(
                data,
                repeatRows=1,
                hAlign="LEFT",
                colWidths=[column_width] * len(table_data.headers),
            )
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E7E6E6")),
                        ("FONTNAME", (0, 0), (-1, 0), "Times-Bold"),
                        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#777777")),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 4),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                        ("TOPPADDING", (0, 0), (-1, -1), 3),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                    ]
                )
            )
            story.append(table)
            story.append(Spacer(1, 6))

    if options.include_disclaimer and document_model.disclaimer:
        story.append(_pdf_para(DISCLAIMER_HEADING, heading_style))
        disclaimer_style = ParagraphStyle("Disclaimer", parent=small_style, fontName="Times-Italic")
        story.append(_pdf_para(document_model.disclaimer, disclaimer_style))

    try:
        doc.build(story)
    except Exception as exc:
        raise DocumentExportGenerationError(f"PDF generation failed: {exc}") from exc
    return output.getvalue()


def generate_export(request: DocumentExportRequest) -> tuple[bytes, GeneratedDocument]:
    document, warnings = _canonicalize(request)
    if request.output_format == ExportFormat.DOCX:
        content = _generate_docx(document, request.options)
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        page_count = None
    else:
        content = _generate_pdf(document, request.options)
        media_type = "application/pdf"
        try:
            with fitz.open(stream=content, filetype="pdf") as pdf:
                page_count = pdf.page_count
        except Exception as exc:
            raise DocumentExportGenerationError(f"generated PDF could not be validated: {exc}") from exc

    if not content:
        raise DocumentExportGenerationError("generated document is empty")
    if len(content) > MAX_OUTPUT_BYTES:
        raise DocumentExportGenerationError(
            f"generated document exceeds the {MAX_OUTPUT_BYTES:,}-byte safety limit"
        )

    extension = request.output_format.value
    filename = request.options.filename or _safe_filename(document.title, extension)
    if not filename.lower().endswith(f".{extension}"):
        filename = _safe_filename(filename, extension)
    result = GeneratedDocument(
        filename=filename,
        media_type=media_type,
        sha256=hashlib.sha256(content).hexdigest(),
        size_bytes=len(content),
        page_count=page_count,
        warnings=warnings,
    )
    return content, result
